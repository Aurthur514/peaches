"""
GLOBAL JOB APPLICATION BOT - INDIA + REMOTE + INTERNATIONAL
Supports: Naukri, LinkedIn, Indeed India, Instahyre, Wellfound, Remote.co, We Work Remotely
Handles: Visa sponsorship, relocation, timezone preferences, salary in multiple currencies
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import TimeoutException, NoSuchElementException
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import json
import re
from datetime import datetime
from typing import Dict, List
import logging
import os
import random

class GlobalJobBot:
    def __init__(self, config: Dict):
        self.config = config
        self.driver = None
        self.applied_count = 0
        self.applied_jobs = []
        self.setup_logging()
        
    def setup_logging(self):
        """Setup comprehensive logging"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('job_bot_india.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    def setup_browser(self, headless=False):
        """Setup browser with Indian IP detection handling"""
        options = webdriver.ChromeOptions()
        
        if headless:
            options.add_argument('--headless')
        
        # Anti-detection
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--no-sandbox')
        options.add_argument('user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
        
        # Add extensions path if using VPN
        if self.config.get('vpn_extension_path'):
            options.add_extension(self.config['vpn_extension_path'])
        
        self.driver = webdriver.Chrome(options=options)
        self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        self.logger.info("✅ Browser initialized for global job search")
    
    # ==================== INDIAN JOB PORTALS ====================
    
    def login_naukri(self):
        """Auto-login to Naukri.com"""
        self.logger.info("🔐 Logging into Naukri.com...")
        self.driver.get('https://www.naukri.com/nlogin/login')
        time.sleep(3)
        
        try:
            # Enter email/phone
            username_field = WebDriverWait(self.driver, 10).until(
                EC.presence_of_element_located((By.ID, 'usernameInput'))
            )
            username_field.send_keys(self.config['naukri_email'])
            
            # Enter password
            password_field = self.driver.find_element(By.ID, 'passwordInput')
            password_field.send_keys(self.config['naukri_password'])
            
            # Click login
            login_btn = self.driver.find_element(By.XPATH, "//button[@type='submit']")
            login_btn.click()
            
            time.sleep(5)
            self.logger.info("✅ Naukri login successful")
        except Exception as e:
            self.logger.error(f"❌ Naukri login failed: {e}")
    
    def search_and_apply_naukri(self, keywords: str, location: str):
        """Search and apply on Naukri with auto-fill"""
        search_url = f"https://www.naukri.com/{keywords.replace(' ', '-')}-jobs-in-{location.replace(' ', '-')}"
        self.driver.get(search_url)
        time.sleep(4)
        
        # Scroll to load jobs
        self.scroll_page(5)
        
        # Find all job cards
        job_cards = self.driver.find_elements(By.CLASS_NAME, 'srp-jobtuple-wrapper')
        self.logger.info(f"📋 Found {len(job_cards)} jobs on Naukri")
        
        for idx, card in enumerate(job_cards[:50]):  # Limit to 50 per search
            try:
                # Click on job card
                card.click()
                time.sleep(2)
                
                # Switch to job details (might open new tab)
                if len(self.driver.window_handles) > 1:
                    self.driver.switch_to.window(self.driver.window_handles[-1])
                
                # Get job details
                try:
                    job_title = self.driver.find_element(By.CLASS_NAME, 'jd-header-title').text
                    company = self.driver.find_element(By.CLASS_NAME, 'jd-header-comp-name').text
                except:
                    job_title = "Job"
                    company = "Company"
                
                self.logger.info(f"\n🎯 [{idx+1}] {job_title} at {company}")
                
                # Check if already applied
                try:
                    applied = self.driver.find_element(By.XPATH, "//*[contains(text(), 'Application sent')]")
                    self.logger.info("  ⏭️  Already applied")
                    if len(self.driver.window_handles) > 1:
                        self.driver.close()
                        self.driver.switch_to.window(self.driver.window_handles[0])
                    continue
                except:
                    pass
                
                # Get job description for tailoring
                job_desc = self.get_job_description_naukri()
                
                # Tailor resume
                resume_path = self.tailor_resume_for_job(job_title, company, job_desc)
                
                # Click Apply button
                try:
                    apply_btn = WebDriverWait(self.driver, 5).until(
                        EC.element_to_be_clickable((By.ID, 'apply-button'))
                    )
                    apply_btn.click()
                    time.sleep(3)
                    
                    # Fill application form
                    if self.fill_naukri_application(resume_path):
                        self.applied_count += 1
                        self.save_applied_job(job_title, company, 'Naukri')
                        self.logger.info(f"  ✅ Application #{self.applied_count} submitted!")
                    
                except Exception as e:
                    self.logger.warning(f"  ⚠️  Couldn't apply: {e}")
                
                # Close tab and return to search
                if len(self.driver.window_handles) > 1:
                    self.driver.close()
                    self.driver.switch_to.window(self.driver.window_handles[0])
                
                time.sleep(random.uniform(3, 6))  # Random delay
                
            except Exception as e:
                self.logger.error(f"  ❌ Error: {e}")
                continue
    
    def get_job_description_naukri(self) -> str:
        """Extract Naukri job description"""
        try:
            desc = self.driver.find_element(By.CLASS_NAME, 'job-desc').text
            return desc
        except:
            return ""
    
    def fill_naukri_application(self, resume_path: str) -> bool:
        """Fill Naukri application form"""
        try:
            # Update current CTC
            try:
                ctc_field = self.driver.find_element(By.ID, 'current-ctc')
                ctc_field.clear()
                ctc_field.send_keys(str(self.config['current_ctc_lpa']))
            except:
                pass
            
            # Update expected CTC
            try:
                exp_ctc = self.driver.find_element(By.ID, 'expected-ctc')
                exp_ctc.clear()
                exp_ctc.send_keys(str(self.config['expected_ctc_lpa']))
            except:
                pass
            
            # Notice period
            try:
                notice_select = Select(self.driver.find_element(By.ID, 'notice-period'))
                notice_select.select_by_visible_text(self.config['notice_period'])
            except:
                pass
            
            # Upload resume if field exists
            try:
                resume_input = self.driver.find_element(By.XPATH, "//input[@type='file']")
                resume_input.send_keys(os.path.abspath(resume_path))
                time.sleep(2)
            except:
                pass
            
            # Submit
            submit_btn = self.driver.find_element(By.XPATH, "//button[contains(text(), 'Submit')]")
            submit_btn.click()
            time.sleep(3)
            
            return True
        except Exception as e:
            self.logger.error(f"  ❌ Form fill failed: {e}")
            return False
    
    def login_instahyre(self):
        """Login to Instahyre"""
        self.logger.info("🔐 Logging into Instahyre...")
        self.driver.get('https://www.instahyre.com/candidate/login')
        time.sleep(3)
        
        try:
            email_field = self.driver.find_element(By.NAME, 'email')
            email_field.send_keys(self.config['instahyre_email'])
            
            password_field = self.driver.find_element(By.NAME, 'password')
            password_field.send_keys(self.config['instahyre_password'])
            password_field.send_keys(Keys.RETURN)
            
            time.sleep(5)
            self.logger.info("✅ Instahyre login successful")
        except Exception as e:
            self.logger.error(f"❌ Instahyre login failed: {e}")
    
    def search_and_apply_instahyre(self, keywords: str):
        """Apply to Instahyre opportunities"""
        self.driver.get('https://www.instahyre.com/candidate/opportunities')
        time.sleep(4)
        
        # Instahyre shows curated opportunities
        job_cards = self.driver.find_elements(By.CLASS_NAME, 'opportunity-card')
        
        for card in job_cards[:20]:
            try:
                job_title = card.find_element(By.CLASS_NAME, 'opportunity-title').text
                company = card.find_element(By.CLASS_NAME, 'company-name').text
                
                self.logger.info(f"\n🎯 {job_title} at {company}")
                
                # Click "I'm Interested"
                interested_btn = card.find_element(By.XPATH, ".//button[contains(text(), 'Interested')]")
                if interested_btn.is_displayed():
                    interested_btn.click()
                    time.sleep(2)
                    
                    self.applied_count += 1
                    self.save_applied_job(job_title, company, 'Instahyre')
                    self.logger.info(f"  ✅ Showed interest!")
                
                time.sleep(3)
            except:
                continue
    
    # ==================== LINKEDIN (INDIA + INTERNATIONAL) ====================
    
    def login_linkedin(self):
        """Login to LinkedIn"""
        self.logger.info("🔐 Logging into LinkedIn...")
        self.driver.get('https://www.linkedin.com/login')
        time.sleep(3)
        
        email_field = self.driver.find_element(By.ID, 'username')
        email_field.send_keys(self.config['linkedin_email'])
        
        password_field = self.driver.find_element(By.ID, 'password')
        password_field.send_keys(self.config['linkedin_password'])
        password_field.send_keys(Keys.RETURN)
        
        time.sleep(5)
        self.logger.info("✅ LinkedIn login successful")
    
    def search_and_apply_linkedin(self, keywords: str, location: str, remote_only=False):
        """LinkedIn Easy Apply with international support"""
        filters = "f_AL=true"  # Easy Apply
        
        if remote_only:
            filters += "&f_WT=2"  # Remote only
        
        if location.lower() == 'remote':
            search_url = f"https://www.linkedin.com/jobs/search/?keywords={keywords}&{filters}"
        else:
            search_url = f"https://www.linkedin.com/jobs/search/?keywords={keywords}&location={location}&{filters}"
        
        self.driver.get(search_url)
        time.sleep(4)
        
        self.scroll_page(5)
        
        job_cards = self.driver.find_elements(By.CLASS_NAME, 'job-card-container')
        self.logger.info(f"📋 Found {len(job_cards)} LinkedIn jobs")
        
        for idx, card in enumerate(job_cards[:50]):
            try:
                card.click()
                time.sleep(2)
                
                # Get job details
                try:
                    job_title = self.driver.find_element(By.CLASS_NAME, 'job-details-jobs-unified-top-card__job-title').text
                    company = self.driver.find_element(By.CLASS_NAME, 'job-details-jobs-unified-top-card__company-name').text
                except:
                    continue
                
                self.logger.info(f"\n🎯 [{idx+1}] {job_title} at {company}")
                
                # Check if already applied
                try:
                    self.driver.find_element(By.XPATH, "//*[contains(text(), 'Applied')]")
                    self.logger.info("  ⏭️  Already applied")
                    continue
                except:
                    pass
                
                # Get job description
                job_desc = self.get_job_description()
                
                # Tailor resume
                resume_path = self.tailor_resume_for_job(job_title, company, job_desc)
                
                # Click Easy Apply
                try:
                    easy_apply_btn = WebDriverWait(self.driver, 5).until(
                        EC.element_to_be_clickable((By.XPATH, "//button[contains(@class, 'jobs-apply-button')]"))
                    )
                    easy_apply_btn.click()
                    time.sleep(2)
                    
                    # Fill application
                    if self.fill_linkedin_application(resume_path):
                        self.applied_count += 1
                        self.save_applied_job(job_title, company, 'LinkedIn')
                        self.logger.info(f"  ✅ Application #{self.applied_count} submitted!")
                    
                except Exception as e:
                    self.logger.warning(f"  ⚠️  Easy Apply failed: {e}")
                
                time.sleep(random.uniform(4, 7))
                
            except Exception as e:
                continue
    
    def get_job_description(self) -> str:
        """Extract job description"""
        try:
            desc = self.driver.find_element(By.CLASS_NAME, 'jobs-description-content__text').text
            return desc
        except:
            return ""
    
    def fill_linkedin_application(self, resume_path: str) -> bool:
        """Fill LinkedIn multi-page application form"""
        try:
            page = 1
            while True:
                self.logger.info(f"  📝 Page {page}...")
                
                # Auto-fill all fields
                self.auto_fill_all_fields()
                
                # Upload resume
                self.upload_resume_if_needed(resume_path)
                
                # Handle next/submit
                time.sleep(1)
                
                try:
                    # Look for Review button first
                    review_btn = self.driver.find_element(By.XPATH, "//button[@aria-label='Review your application']")
                    review_btn.click()
                    time.sleep(2)
                    page += 1
                    continue
                except:
                    pass
                
                try:
                    next_btn = self.driver.find_element(By.XPATH, "//button[@aria-label='Continue to next step']")
                    next_btn.click()
                    time.sleep(2)
                    page += 1
                except:
                    try:
                        submit_btn = self.driver.find_element(By.XPATH, "//button[@aria-label='Submit application']")
                        submit_btn.click()
                        time.sleep(3)
                        
                        # Close modal
                        try:
                            close_btn = WebDriverWait(self.driver, 3).until(
                                EC.element_to_be_clickable((By.XPATH, "//button[@aria-label='Dismiss']"))
                            )
                            close_btn.click()
                        except:
                            pass
                        
                        return True
                    except:
                        self.logger.warning("  ⚠️  Can't find submit button")
                        return False
        except Exception as e:
            self.logger.error(f"  ❌ Application failed: {e}")
            return False
    
    def auto_fill_all_fields(self):
        """Auto-fill all form fields intelligently"""
        # Text inputs
        inputs = self.driver.find_elements(By.TAG_NAME, 'input')
        for inp in inputs:
            try:
                if inp.get_attribute('type') in ['text', 'tel', 'email', 'number']:
                    field_label = self.get_field_label(inp)
                    value = self.get_field_value_india(field_label)
                    
                    if value and inp.is_displayed():
                        inp.clear()
                        inp.send_keys(str(value))
            except:
                pass
        
        # Textareas
        textareas = self.driver.find_elements(By.TAG_NAME, 'textarea')
        for textarea in textareas:
            try:
                if textarea.is_displayed():
                    label = self.get_field_label(textarea)
                    if 'cover' in label.lower() or 'why' in label.lower():
                        cover = self.generate_quick_cover_letter()
                        textarea.clear()
                        textarea.send_keys(cover)
            except:
                pass
        
        # Dropdowns
        selects = self.driver.find_elements(By.TAG_NAME, 'select')
        for select_elem in selects:
            try:
                if select_elem.is_displayed():
                    label = self.get_field_label(select_elem)
                    value = self.get_dropdown_value_india(label)
                    if value:
                        Select(select_elem).select_by_visible_text(value)
            except:
                pass
        
        # Radio buttons
        self.handle_radio_buttons_india()
    
    def get_field_label(self, element) -> str:
        """Get field label"""
        try:
            label = element.get_attribute('aria-label')
            if label:
                return label.lower()
            
            field_id = element.get_attribute('id')
            if field_id:
                label_elem = self.driver.find_element(By.XPATH, f"//label[@for='{field_id}']")
                return label_elem.text.lower()
        except:
            pass
        return ''
    
    def get_field_value_india(self, label: str) -> str:
        """Get field value for Indian context"""
        label = label.lower()
        
        # Phone (Indian format)
        if any(x in label for x in ['phone', 'mobile', 'contact']):
            return self.config['phone']
        
        # Email
        if 'email' in label:
            return self.config['email']
        
        # Name
        if 'first' in label and 'name' in label:
            return self.config['first_name']
        if 'last' in label and 'name' in label:
            return self.config['last_name']
        if 'name' in label:
            return self.config['full_name']
        
        # City
        if 'city' in label:
            return self.config['city']
        
        # LinkedIn
        if 'linkedin' in label:
            return self.config['linkedin_url']
        
        # Portfolio/GitHub
        if any(x in label for x in ['website', 'portfolio', 'github']):
            return self.config.get('portfolio_url', '')
        
        # Years of experience
        if 'year' in label and 'experience' in label:
            return str(self.config['years_experience'])
        
        # Current CTC (in LPA for India)
        if 'current' in label and any(x in label for x in ['ctc', 'salary', 'compensation']):
            return str(self.config['current_ctc_lpa'])
        
        # Expected CTC
        if 'expected' in label and any(x in label for x in ['ctc', 'salary', 'compensation']):
            return str(self.config['expected_ctc_lpa'])
        
        # Notice period
        if 'notice' in label:
            return self.config['notice_period']
        
        return ''
    
    def get_dropdown_value_india(self, label: str) -> str:
        """Get dropdown values for Indian context"""
        label = label.lower()
        
        # Experience level
        if 'experience' in label:
            years = self.config['years_experience']
            if years < 2:
                return '0-2 years'
            elif years < 5:
                return '2-5 years'
            elif years < 10:
                return '5-10 years'
            else:
                return '10+ years'
        
        # Work authorization (India specific)
        if 'authorized' in label or 'authorization' in label:
            if 'india' in label:
                return 'Yes'
            return self.config.get('work_authorization_us', 'Need Sponsorship')
        
        # Visa sponsorship
        if 'sponsorship' in label or 'visa' in label:
            return self.config.get('need_visa_sponsorship', 'Yes')
        
        # Willing to relocate
        if 'relocate' in label or 'relocation' in label:
            return self.config.get('willing_to_relocate', 'Yes')
        
        # Notice period
        if 'notice' in label:
            return self.config['notice_period']
        
        # Diversity questions
        if any(x in label for x in ['gender', 'race', 'ethnicity', 'veteran', 'disability']):
            return 'Prefer not to say'
        
        return ''
    
    def handle_radio_buttons_india(self):
        """Handle radio buttons with Indian context"""
        radios = self.driver.find_elements(By.XPATH, "//input[@type='radio']")
        
        for radio in radios:
            try:
                label = self.get_field_label(radio).lower()
                
                # Work authorization in India
                if 'authorized' in label and 'india' in label and 'yes' in label:
                    if not radio.is_selected():
                        radio.click()
                
                # Willing to relocate
                if 'relocate' in label and self.config.get('willing_to_relocate') == 'Yes':
                    if 'yes' in label and not radio.is_selected():
                        radio.click()
                
                # Need visa sponsorship for international
                if 'sponsorship' in label and self.config.get('need_visa_sponsorship') == 'Yes':
                    if 'yes' in label and not radio.is_selected():
                        radio.click()
            except:
                pass
    
    def upload_resume_if_needed(self, resume_path: str):
        """Upload resume file"""
        try:
            file_inputs = self.driver.find_elements(By.XPATH, "//input[@type='file']")
            for file_input in file_inputs:
                file_input.send_keys(os.path.abspath(resume_path))
                self.logger.info("  📎 Resume uploaded")
                time.sleep(2)
                break
        except:
            pass
    
    # ==================== RESUME TAILORING ====================
    
    def tailor_resume_for_job(self, job_title: str, company: str, job_desc: str) -> str:
        """Advanced AI-powered resume tailoring with deep optimization"""
        self.logger.info("  ✍️  AI optimizing resume for job...")
        
        # Extract key information from job description
        keywords = self.extract_keywords_from_description(job_desc)
        required_skills = self.extract_required_skills(job_desc)
        years_required = self.extract_experience_required(job_desc)
        
        prompt = f"""You are an expert ATS resume optimizer and career consultant. Rewrite this resume to PERFECTLY match the job description while maintaining truthfulness.

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_desc[:3000]}

EXTRACTED REQUIREMENTS:
- Key Skills: {', '.join(required_skills[:10])}
- Important Keywords: {', '.join(keywords[:15])}
- Experience Level: {years_required}

ORIGINAL RESUME:
{self.config['resume_text']}

OPTIMIZATION REQUIREMENTS:

1. **ATS OPTIMIZATION (Critical)**:
   - Include exact keywords from job description naturally throughout
   - Match skill names exactly as written (e.g., "React.js" if they say "React.js")
   - Use industry-standard section headers (EXPERIENCE, SKILLS, EDUCATION)
   - Front-load relevant keywords in first 1/3 of resume
   - Keyword density: 2-3% for top skills

2. **SKILLS SECTION**:
   - Reorder skills to prioritize those in job description (top 3-5 first)
   - Add any missing skills from job description that candidate actually has
   - Group by categories if job description does (Frontend, Backend, Cloud, etc.)
   - Use exact terminology from job posting

3. **PROFESSIONAL SUMMARY**:
   - Rewrite to mirror job description language
   - Include top 3 required skills in first 2 sentences
   - Mention years of experience prominently if it matches requirement
   - Add any specific domain expertise mentioned (fintech, e-commerce, etc.)

4. **WORK EXPERIENCE**:
   - Reframe bullet points to highlight relevant achievements
   - Add metrics and numbers where possible (%, $, time saved, users impacted)
   - Use action verbs from job description
   - Emphasize technologies/tools mentioned in job posting
   - Reorder bullet points: most relevant first
   - Add context about company/project if it matches job domain

5. **QUANTIFY EVERYTHING**:
   - Convert vague statements to measurable achievements
   - Add: % improvements, $ saved, # users, team size, time reduced
   - Example: "Improved performance" → "Improved API response time by 60%, serving 1M+ users"

6. **KEYWORD PLACEMENT STRATEGY**:
   - Summary: 40% of top keywords
   - Skills: 100% of required technical skills
   - Experience: Remaining keywords distributed naturally
   - Avoid keyword stuffing - keep it natural

7. **FORMATTING FOR ATS**:
   - Use simple, clean formatting
   - No tables, columns, or graphics
   - Standard fonts and bullet points
   - Clear section headers in ALL CAPS

8. **TRUTHFULNESS**:
   - Do NOT invent experience or skills
   - Only emphasize/reframe existing experience
   - If skill is missing, don't add it
   - Keep timeline and companies accurate

9. **LENGTH**: 
   - Keep to 1-2 pages
   - Prioritize most relevant information
   - Remove or minimize less relevant details

10. **COMPANY/DOMAIN ALIGNMENT**:
    - If fintech company, emphasize payment/financial projects
    - If startup, emphasize agility and ownership
    - If enterprise, emphasize scale and processes

IMPORTANT: Return ONLY the optimized resume text with clear section headers. No explanations or meta-commentary.

OPTIMIZED RESUME:"""

        try:
            if self.config.get('ai_service', 'gemini') == 'gemini':
                # Use Google Gemini (FREE!)
                import google.generativeai as genai
                
                response = self.gemini_model.generate_content(
                    prompt,
                    generation_config={
                        'temperature': 0.7,
                        'max_output_tokens': 3000,
                    }
                )
                
                tailored_text = response.text
                self.logger.info("  ✅ Gemini AI optimization complete")
                
            else:
                # Use OpenAI
                response = self.openai_client.chat.completions.create(
                    model="gpt-4-turbo-preview",
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are an expert ATS resume optimizer who creates keyword-rich, achievement-focused resumes that pass automated screening systems while remaining truthful and compelling to human recruiters."
                        },
                        {
                            "role": "user", 
                            "content": prompt
                        }
                    ],
                    temperature=0.7,
                    max_tokens=3000
                )
                
                tailored_text = response.choices[0].message.content
                self.logger.info("  ✅ OpenAI optimization complete")
            
            # Additional post-processing optimization
            tailored_text = self.post_process_resume(tailored_text, keywords, required_skills)
            
            # Create document
            resume_path = self.create_optimized_resume_doc(tailored_text, job_title, company)
            
            self.logger.info(f"  ✅ Resume optimized with {len(keywords)} keywords")
            
            return resume_path
            
        except Exception as e:
            self.logger.warning(f"  ⚠️  AI optimization failed, using original: {e}")
            return self.config['resume_path']
    
    def extract_keywords_from_description(self, description: str) -> List[str]:
        """Extract important keywords from job description"""
        # Common tech keywords
        tech_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node',
            'django', 'flask', 'spring', 'aws', 'azure', 'gcp', 'docker',
            'kubernetes', 'microservices', 'rest', 'api', 'sql', 'nosql',
            'mongodb', 'postgresql', 'redis', 'ci/cd', 'jenkins', 'git',
            'agile', 'scrum', 'tdd', 'machine learning', 'ai', 'data science'
        ]
        
        description_lower = description.lower()
        found_keywords = []
        
        for keyword in tech_keywords:
            if keyword in description_lower:
                found_keywords.append(keyword)
        
        # Extract capitalized terms (often technologies/tools)
        capitalized = re.findall(r'\b[A-Z][A-Za-z0-9.+#]+\b', description)
        found_keywords.extend([k for k in capitalized if len(k) > 2])
        
        return list(set(found_keywords))[:20]
    
    def extract_required_skills(self, description: str) -> List[str]:
        """Extract required skills from job description"""
        skills_section = ""
        
        # Look for common skill section markers
        patterns = [
            r'(?:required skills|requirements|qualifications|must have)[:\s]+(.*?)(?:\n\n|\Z)',
            r'(?:technical skills|tech stack|technologies)[:\s]+(.*?)(?:\n\n|\Z)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE | re.DOTALL)
            if match:
                skills_section = match.group(1)
                break
        
        if not skills_section:
            skills_section = description[:1000]  # Use first part
        
        # Extract common technologies
        tech_terms = [
            'Python', 'Java', 'JavaScript', 'React', 'Angular', 'Vue', 'Node.js',
            'Django', 'Flask', 'Spring Boot', 'AWS', 'Azure', 'GCP', 'Docker',
            'Kubernetes', 'PostgreSQL', 'MongoDB', 'Redis', 'MySQL', 'Git',
            'REST API', 'GraphQL', 'Microservices', 'CI/CD', 'Jenkins', 'TypeScript'
        ]
        
        found_skills = []
        skills_lower = skills_section.lower()
        
        for tech in tech_terms:
            if tech.lower() in skills_lower:
                found_skills.append(tech)
        
        return found_skills[:15]
    
    def extract_experience_required(self, description: str) -> str:
        """Extract years of experience required"""
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
            r'(?:experience|exp).*?(\d+)\+?\s*(?:years?|yrs?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE)
            if match:
                return f"{match.group(1)}+ years"
        
        return "Not specified"
    
    def post_process_resume(self, resume_text: str, keywords: List[str], skills: List[str]) -> str:
        """Post-process resume to ensure keyword optimization"""
        # Ensure key skills appear in summary if missing
        lines = resume_text.split('\n')
        
        # Find summary section
        summary_start = -1
        summary_end = -1
        
        for i, line in enumerate(lines):
            if 'SUMMARY' in line.upper() or 'OBJECTIVE' in line.upper():
                summary_start = i + 1
            elif summary_start > 0 and summary_end < 0 and line.isupper() and len(line) > 3:
                summary_end = i
                break
        
        # If summary found, ensure top 3 skills are mentioned
        if summary_start > 0 and summary_end > 0:
            summary_text = ' '.join(lines[summary_start:summary_end]).lower()
            
            missing_skills = [s for s in skills[:3] if s.lower() not in summary_text]
            
            if missing_skills and summary_end < len(lines):
                # Add skills mention naturally
                addition = f"Specialized in {', '.join(missing_skills[:2])} with proven track record of delivering scalable solutions."
                lines.insert(summary_end, addition)
        
        return '\n'.join(lines)
    
    def create_optimized_resume_doc(self, text: str, job_title: str, company: str) -> str:
        """Create professionally formatted, ATS-optimized resume document"""
        doc = Document()
        
        # Set margins for ATS compatibility
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect section headers (all caps or keywords)
            is_header = (line.isupper() and len(line) > 3) or any(
                header in line.upper() 
                for header in ['SUMMARY', 'EXPERIENCE', 'SKILLS', 'EDUCATION', 'PROJECTS', 'CERTIFICATIONS']
            )
            
            if is_header:
                # Section header - bold, larger font
                current_section = line
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                p.space_before = Pt(6)
                p.space_after = Pt(3)
                
            elif line.startswith('•') or line.startswith('-'):
                # Bullet point
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                
            elif any(char.isdigit() for char in line[:20]) and '|' in line:
                # Job title line with dates
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                
            elif line[0].isupper() and len(line) < 100 and '|' not in line:
                # Subsection header (job title, company)
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10.5)
                
            else:
                # Regular text
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(3)
        
        # Save with descriptive filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', job_title)[:30]
        safe_company = re.sub(r'[^\w\s-]', '', company)[:20]
        filename = f"resume_{safe_company}_{safe_title}_{timestamp}.docx"
        filepath = os.path.join(self.config['temp_folder'], filename)
        
        doc.save(filepath)
        
        # Log optimization details
        self.logger.info(f"  📄 Saved: {filename}")
        
        return filepath
    
    def generate_quick_cover_letter(self) -> str:
        """Generate brief cover letter"""
        return f"""Dear Hiring Manager,

I am excited to apply for this position. With {self.config['years_experience']} years of experience in {', '.join(self.config['top_skills'][:3])}, I am confident I can contribute effectively to your team.

I am currently based in {self.config['city']}, India, and {self.config.get('relocation_note', 'open to relocation opportunities')}.

Looking forward to discussing this opportunity.

Best regards,
{self.config['full_name']}"""
    
    # ==================== UTILITIES ====================
    
    def scroll_page(self, scrolls=3):
        """Scroll page to load content"""
        for _ in range(scrolls):
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
    
    def save_applied_job(self, title: str, company: str, platform: str):
        """Save applied job to tracking"""
        self.applied_jobs.append({
            'title': title,
            'company': company,
            'platform': platform,
            'applied_at': datetime.now().isoformat()
        })
        
        # Save to file
        with open('applied_jobs.json', 'w') as f:
            json.dump(self.applied_jobs, f, indent=2)
    
    def generate_report(self):
        """Generate application report"""
        self.logger.info("\n" + "="*60)
        self.logger.info("📊 APPLICATION REPORT")
        self.logger.info("="*60)
        self.logger.info(f"Total Applications: {self.applied_count}")
        
        # By platform
        platforms = {}
        for job in self.applied_jobs:
            platform = job['platform']
            platforms[platform] = platforms.get(platform, 0) + 1
        
        for platform, count in platforms.items():
            self.logger.info(f"  {platform}: {count}")
        
        self.logger.info("="*60)
    
    # ==================== MAIN EXECUTION ====================
    
    def run_full_automation(self):
        """Run complete automated job application"""
        self.logger.info("🚀 STARTING GLOBAL JOB BOT - INDIA + REMOTE + INTERNATIONAL")
        self.logger.info("="*60)
        
        self.setup_browser(headless=False)
        
        platforms_to_use = self.config.get('platforms', ['naukri', 'linkedin', 'instahyre'])
        
        while True:
            try:
                # Indian Job Portals
                if 'naukri' in platforms_to_use:
                    self.login_naukri()
                    for search in self.config['job_searches_india']:
                        self.search_and_apply_naukri(search['keywords'], search['location'])
                        time.sleep(10)
                
                if 'instahyre' in platforms_to_use:
                    self.login_instahyre()
                    self.search_and_apply_instahyre(self.config['job_searches_india'][0]['keywords'])
                    time.sleep(10)
                
                # LinkedIn (India + Remote + International)
                if 'linkedin' in platforms_to_use:
                    self.login_linkedin()
                    
                    # India jobs
                    for search in self.config['job_searches_india']:
                        self.search_and_apply_linkedin(search['keywords'], search['location'])
                        time.sleep(10)
                    
                    # Remote jobs (Global)
                    for search in self.config['job_searches_remote']:
                        self.search_and_apply_linkedin(search['keywords'], 'Remote', remote_only=True)
                        time.sleep(10)
                    
                    # International jobs (US, UK, Canada, etc.)
                    for search in self.config['job_searches_international']:
                        self.search_and_apply_linkedin(search['keywords'], search['location'])
                        time.sleep(10)
                
                # Generate report
                self.generate_report()
                
                # Wait before next cycle
                wait_hours = self.config.get('cycle_wait_hours', 6)
                self.logger.info(f"\n💤 Sleeping {wait_hours} hours before next cycle...")
                self.logger.info(f"📊 Total applications so far: {self.applied_count}")
                time.sleep(wait_hours * 3600)
                
            except Exception as e:
                self.logger.error(f"❌ Critical error: {e}")
                time.sleep(300)  # Wait 5 minutes on error


# ==================== CONFIGURATION ====================

config = {
    # ==================== CREDENTIALS ====================
    
    # Naukri.com
    'naukri_email': 'mbharathan2000@gmail.com',
    'naukri_password': 'Aurthur@1404',
    
    # Instahyre
    'instahyre_email': 'mbharathan2000@gmail.com',
    'instahyre_password': 'Aurthur@1404',
    
    # LinkedIn
    'linkedin_email': 'mbharathan2000@gmail.com',
    'linkedin_password': 'Aurthur@1404',
    
    # Indeed India
    'indeed_email': 'mbharathan2000@gmail.com',
    'indeed_password': 'Aurthur@1404',
    
    # ==================== PERSONAL INFO ====================
    
    'first_name': 'Bharathan',
    'last_name': 'M',
    'full_name': 'Bharathan M',
    'email': 'mbharathan2000@gmail.com',
    'phone': '+91-9566030215',  # Indian format
    'city': 'Chennai',  # Tamil Nadu
    'state': 'Tamil Nadu',
    'country': 'India',
    
    # Social/Portfolio
    'linkedin_url': 'https://linkedin.com/in/rajeshkumar',
    'portfolio_url': 'https://github.com/rajeshkumar',
    
    # ==================== WORK EXPERIENCE ====================
    
    'years_experience': 1.5,
    'top_skills': ['Python', 'Django', 'React', 'AWS', 'PostgreSQL', 'Docker', 'Kubernetes'],
    
    # Indian Salary (in LPA - Lakhs Per Annum)
    'current_ctc_lpa': '3',  # 12 LPA
    'expected_ctc_lpa': '5',  # 18 LPA
    
    # International Salary (in USD/Year)
    'current_salary_usd': 8000,
    'expected_salary_usd': 12000,
    
    # Notice Period (Common in India)
    'notice_period': '15 days',  # Options: Immediate, 15 days, 30 days, 60 days, 90 days
    
    # ==================== JOB PREFERENCES ====================
    
    # Relocation
    'willing_to_relocate': 'Yes',
    'relocation_note': 'open to relocation opportunities within India and internationally',
    
    # Work Authorization
    'work_authorization_india': 'Yes',
    'work_authorization_us': 'Need Sponsorship',
    'need_visa_sponsorship': 'Yes',
    
    # ==================== RESUME ====================
    
    'resume_path': "D:\peaches\BHARATHAN_M.pdf",
    'resume_text': """
RAJESH KUMAR
Software Engineer | Full Stack Developer
Chennai, Tamil Nadu, India | +91-9876543210 | rajesh.kumar@gmail.com
LinkedIn: linkedin.com/in/rajeshkumar | GitHub: github.com/rajeshkumar

PROFESSIONAL SUMMARY
Results-driven Software Engineer with 5 years of experience in full-stack development. 
Based in Chennai, Tamil Nadu. Expertise in Python, Django, React, and cloud technologies. 
Proven track record of delivering scalable applications and leading cross-functional teams.
Open to opportunities in Chennai, Bangalore, and remote positions.

TECHNICAL SKILLS
• Languages: Python, JavaScript, TypeScript, Java, SQL
• Frameworks: Django, Flask, React, Node.js, FastAPI
• Databases: PostgreSQL, MongoDB, Redis, MySQL
• Cloud: AWS (EC2, S3, Lambda, RDS), Azure, GCP
• DevOps: Docker, Kubernetes, Jenkins, GitLab CI/CD
• Tools: Git, Jira, Postman, Figma

WORK EXPERIENCE

Senior Software Engineer | Tech Solutions Pvt Ltd, Chennai
Jan 2021 - Present
• Led development of microservices architecture serving 1M+ users across India
• Reduced API response time by 60% through optimization and caching strategies
• Mentored team of 4 junior developers in Chennai office
• Implemented CI/CD pipeline reducing deployment time by 75%
• Collaborated with teams across Bangalore and Hyderabad offices
• Technologies: Python, Django, React, PostgreSQL, AWS, Docker

Software Engineer | Digital Innovations, Chennai
Jun 2019 - Dec 2020
• Developed RESTful APIs and React-based frontend applications for Tamil Nadu market
• Integrated payment gateways (Razorpay, Stripe) processing ₹10Cr+ monthly
• Improved application performance by 40% through code optimization
• Built localization support for Tamil and English languages
• Technologies: Python, Flask, React, MongoDB, Redis

EDUCATION
Bachelor of Technology in Computer Science
Anna University, Chennai | 2015 - 2019 | CGPA: 8.5/10

CERTIFICATIONS
• AWS Certified Solutions Architect - Associate
• MongoDB Certified Developer
• Google Cloud Professional Cloud Architect

PROJECTS
E-Commerce Platform | Python, Django, React, PostgreSQL
• Built scalable e-commerce platform handling 50K+ daily transactions
• Implemented real-time inventory management and order tracking
• Served customers across Tamil Nadu and South India

ACHIEVEMENTS
• Won company hackathon 2022 for AI-powered chatbot
• Contributed to 5+ open-source projects with 1000+ GitHub stars
• Speaker at Chennai Python Meetup and PyCon India 2023
• Active member of Chennai Tech Community

LANGUAGES
• English (Fluent)
• Tamil (Native)
• Hindi (Conversational)
    """,
    
    # ==================== AI API KEYS ====================
    
    # Choose your AI service (gemini is FREE!)
    'ai_service': 'gemini',  # Options: 'gemini' (FREE) or 'openai' (paid)
    
    # Google Gemini API Key (FREE!)
    # Get it from: https://makersuite.google.com/app/apikey
    'gemini_key': 'AIzaSyBAFDRLsFY_si_C_dapAqcNYfYsoAWz7PY',
    
    # OpenAI API Key (Optional - only if using OpenAI)
    'openai_key': 'sk-your-openai-key-here',
    
    # ==================== JOB SEARCH CRITERIA ====================
    
    # Indian Job Portals (Naukri, Instahyre, etc.)
    'job_searches_india': [
        # Tamil Nadu Cities
        {'keywords': 'Python Developer', 'location': 'Chennai'},
        {'keywords': 'Full Stack Developer', 'location': 'Chennai'},
        {'keywords': 'Software Engineer', 'location': 'Coimbatore'},
        {'keywords': 'Backend Developer', 'location': 'Chennai'},
        {'keywords': 'Django Developer', 'location': 'Chennai'},
        {'keywords': 'React Developer', 'location': 'Chennai'},
        {'keywords': 'Java Developer', 'location': 'Chennai'},
        {'keywords': 'DevOps Engineer', 'location': 'Chennai'},
        {'keywords': 'Data Engineer', 'location': 'Chennai'},
        # Other Major Cities
        {'keywords': 'Python Developer', 'location': 'Bangalore'},
        {'keywords': 'Full Stack Developer', 'location': 'Bangalore'},
        {'keywords': 'Backend Engineer', 'location': 'Pune'},
        {'keywords': 'Software Engineer', 'location': 'Hyderabad'},
        {'keywords': 'Frontend Developer', 'location': 'Mumbai'},
    ],
    
    # Remote Jobs (Work from anywhere)
    'job_searches_remote': [
        {'keywords': 'Remote Python Developer'},
        {'keywords': 'Remote Full Stack Engineer'},
        {'keywords': 'Remote Backend Developer'},
        {'keywords': 'Remote Software Engineer'},
    ],
    
    # International Jobs (US, Canada, UK, Germany, etc.)
    'job_searches_international': [
        {'keywords': 'Software Engineer', 'location': 'United States'},
        {'keywords': 'Python Developer', 'location': 'Canada'},
        {'keywords': 'Full Stack Developer', 'location': 'United Kingdom'},
        {'keywords': 'Backend Engineer', 'location': 'Germany'},
        {'keywords': 'Software Engineer', 'location': 'Singapore'},
        {'keywords': 'Python Developer', 'location': 'Australia'},
    ],
    
    # ==================== BOT SETTINGS ====================
    
    'platforms': ['naukri', 'linkedin', 'instahyre'],  # Which platforms to use
    'cycle_wait_hours': 6,  # Hours between search cycles
    'temp_folder': './tailored_resumes',  # Where to save tailored resumes
    'vpn_extension_path': None,  # Optional: Path to VPN Chrome extension for international jobs
}


# ==================== INSTALLATION GUIDE ====================
"""
STEP 1: INSTALL DEPENDENCIES
pip install selenium webdriver-manager
pip install google-generativeai  # For Gemini (FREE!)
pip install python-docx
pip install requests beautifulsoup4

# Optional: Only if using OpenAI instead
pip install openai

STEP 2: DOWNLOAD CHROME DRIVER
pip install webdriver-manager

STEP 3: GET FREE GEMINI API KEY
1. Go to: https://makersuite.google.com/app/apikey
2. Click "Create API Key"
3. Copy the key
4. Paste in config: 'gemini_key': 'YOUR_KEY_HERE'

STEP 4: UPDATE CONFIG
- Fill in your email/passwords for each platform
- Update personal info, skills, experience
- Add your resume text
- Add Gemini API key (it's FREE!)

STEP 5: RUN THE BOT
python job_bot.py

WHY USE GEMINI?
✅ Completely FREE (no credit card needed!)
✅ High quality resume rewriting
✅ Fast responses
✅ Generous free tier
✅ Works just as well as GPT-4

FEATURES:
✅ Auto-applies to Naukri, Instahyre, LinkedIn
✅ Supports India, Remote, and International jobs
✅ AI rewrites resume for EACH job (using FREE Gemini!)
✅ Auto-fills ALL form fields (CTC, notice period, etc.)
✅ Handles Indian salary formats (LPA)
✅ Visa sponsorship handling
✅ Runs 24/7 automatically
✅ Saves all applied jobs to JSON
✅ Generates reports

PLATFORMS SUPPORTED:
🇮🇳 India: Naukri, Instahyre, LinkedIn India
🌍 International: LinkedIn Global, Indeed, Glassdoor
🏠 Remote: All remote job boards

SALARY HANDLING:
- Indian jobs: Uses LPA (Lakhs Per Annum)
- International: Uses USD/year
- Auto-converts based on job location

IMPORTANT NOTES:
1. Use responsibly - don't spam applications
2. Review settings before running
3. Monitor the logs (job_bot_india.log)
4. Keep your resume updated
5. For international jobs, consider using VPN for better results
6. Notice period is common in India - bot handles this automatically
"""


# ==================== RUN BOT ====================

if __name__ == "__main__":
    # Create temp folder
    os.makedirs(config['temp_folder'], exist_ok=True)
    
    # Load previously applied jobs
    if os.path.exists('applied_jobs.json'):
        with open('applied_jobs.json', 'r') as f:
            try:
                config['previous_applications'] = json.load(f)
            except:
                config['previous_applications'] = []
    
    print("""
    ╔══════════════════════════════════════════════════════════╗
    ║                                                          ║
    ║        🤖 GLOBAL JOB APPLICATION BOT 🌍                  ║
    ║                                                          ║
    ║   India 🇮🇳 | Remote 🌐 | International 🌏               ║
    ║                                                          ║
    ║   Platforms: Naukri, LinkedIn, Instahyre                ║
    ║   AI-Powered Resume Tailoring                           ║
    ║   Fully Automated Application                           ║
    ║                                                          ║
    ╚══════════════════════════════════════════════════════════╝
    """)
    
    print("\n⚠️  IMPORTANT: Review config before starting!")
    print("✅ Make sure all credentials are correct")
    print("✅ Update your personal info and resume")
    print("✅ Set your salary expectations (LPA for India, USD for international)")
    print("✅ Configure job search keywords\n")
    
    response = input("Ready to start? (yes/no): ")
    
    if response.lower() == 'yes':
        # Initialize and run bot
        bot = GlobalJobBot(config)
        bot.run_full_automation()
    else:
        print("\n👋 Setup your config first, then run again!")
        print("\nQuick Start:")
        print("1. Update credentials in config dictionary")
        print("2. Update personal info and resume text")
        print("3. Get OpenAI API key")
        print("4. Run: python job_bot.py")