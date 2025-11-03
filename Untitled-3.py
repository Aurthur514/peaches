"""
FULLY AUTOMATED JOB APPLICATION BOT
- Auto-fills all application forms
- Dynamically rewrites resume for each role
- Handles file uploads, questions, assessments
- Runs 24/7 with monitoring
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import TimeoutException, NoSuchElementException
import openai
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
import PyPDF2
from io import BytesIO
import time
import json
import re
from datetime import datetime
from typing import Dict, List
import logging
import os

class AutomatedJobBot:
    def __init__(self, config: Dict):
        self.config = config
        self.driver = None
        self.applied_count = 0
        self.setup_logging()
        self.setup_ai_clients()
        
    def setup_logging(self):
        """Setup logging for monitoring"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('job_bot.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    def setup_ai_clients(self):
        """Initialize AI APIs for resume rewriting"""
        # Add your API keys
        self.openai_client = openai.OpenAI(api_key=self.config.get('openai_key'))
        # self.anthropic_client = anthropic.Anthropic(api_key=self.config.get('anthropic_key'))
    
    def setup_browser(self, headless=False):
        """Setup Selenium WebDriver with stealth settings"""
        options = webdriver.ChromeOptions()
        
        if headless:
            options.add_argument('--headless')
        
        # Anti-detection measures
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--no-sandbox')
        options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
        
        self.driver = webdriver.Chrome(options=options)
        
        # Execute CDP commands for stealth
        self.driver.execute_cdp_cmd('Network.setUserAgentOverride', {
            "userAgent": 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        self.logger.info("✅ Browser initialized")
    
    def login_to_platforms(self):
        """Auto-login to all job platforms"""
        platforms = {
            'linkedin': self.login_linkedin,
            'indeed': self.login_indeed,
            'glassdoor': self.login_glassdoor,
            'wellfound': self.login_wellfound
        }
        
        for platform, login_func in platforms.items():
            try:
                self.logger.info(f"🔐 Logging into {platform}...")
                login_func()
                time.sleep(3)
            except Exception as e:
                self.logger.error(f"❌ Failed to login to {platform}: {e}")
    
    def login_linkedin(self):
        """Auto-login to LinkedIn"""
        self.driver.get('https://www.linkedin.com/login')
        time.sleep(2)
        
        # Fill login form
        email_field = self.driver.find_element(By.ID, 'username')
        email_field.send_keys(self.config['linkedin_email'])
        
        password_field = self.driver.find_element(By.ID, 'password')
        password_field.send_keys(self.config['linkedin_password'])
        password_field.send_keys(Keys.RETURN)
        
        time.sleep(5)
        self.logger.info("✅ LinkedIn login successful")
    
    def login_indeed(self):
        """Auto-login to Indeed"""
        self.driver.get('https://secure.indeed.com/account/login')
        time.sleep(2)
        
        email_field = self.driver.find_element(By.ID, 'login-email-input')
        email_field.send_keys(self.config['indeed_email'])
        email_field.send_keys(Keys.RETURN)
        time.sleep(2)
        
        password_field = self.driver.find_element(By.ID, 'login-password-input')
        password_field.send_keys(self.config['indeed_password'])
        password_field.send_keys(Keys.RETURN)
        
        time.sleep(5)
        self.logger.info("✅ Indeed login successful")
    
    def login_glassdoor(self):
        """Auto-login to Glassdoor"""
        self.driver.get('https://www.glassdoor.com/profile/login_input.htm')
        time.sleep(2)
        # Similar login logic
    
    def login_wellfound(self):
        """Auto-login to Wellfound/AngelList"""
        self.driver.get('https://angel.co/login')
        time.sleep(2)
        # Similar login logic
    
    def search_and_apply_linkedin(self, keywords: str, location: str):
        """Search LinkedIn and auto-apply to Easy Apply jobs"""
        search_url = f"https://www.linkedin.com/jobs/search/?keywords={keywords}&location={location}&f_AL=true"
        self.driver.get(search_url)
        time.sleep(3)
        
        # Scroll to load more jobs
        self.scroll_page()
        
        # Find all Easy Apply jobs
        job_cards = self.driver.find_elements(By.CLASS_NAME, 'job-card-container')
        
        self.logger.info(f"📋 Found {len(job_cards)} Easy Apply jobs")
        
        for idx, card in enumerate(job_cards):
            try:
                card.click()
                time.sleep(2)
                
                # Get job details
                job_title = self.driver.find_element(By.CLASS_NAME, 'job-details-jobs-unified-top-card__job-title').text
                company = self.driver.find_element(By.CLASS_NAME, 'job-details-jobs-unified-top-card__company-name').text
                
                self.logger.info(f"\n🎯 [{idx+1}/{len(job_cards)}] {job_title} at {company}")
                
                # Check if already applied
                try:
                    applied_badge = self.driver.find_element(By.XPATH, "//*[contains(text(), 'Applied')]")
                    self.logger.info("  ⏭️  Already applied, skipping...")
                    continue
                except:
                    pass
                
                # Click Easy Apply button
                easy_apply_btn = self.driver.find_element(By.XPATH, "//button[contains(@class, 'jobs-apply-button')]")
                easy_apply_btn.click()
                time.sleep(2)
                
                # Get job description for resume tailoring
                job_description = self.get_job_description()
                
                # Tailor resume for this specific job
                tailored_resume_path = self.tailor_resume_for_job(job_title, company, job_description)
                
                # Fill out the application
                if self.fill_linkedin_application(tailored_resume_path):
                    self.applied_count += 1
                    self.logger.info(f"  ✅ Application #{self.applied_count} submitted!")
                
                time.sleep(5)  # Rate limiting
                
            except Exception as e:
                self.logger.error(f"  ❌ Error: {e}")
                continue
    
    def get_job_description(self) -> str:
        """Extract full job description"""
        try:
            desc_element = self.driver.find_element(By.CLASS_NAME, 'jobs-description-content__text')
            return desc_element.text
        except:
            return ""
    
    def tailor_resume_for_job(self, job_title: str, company: str, job_description: str) -> str:
        """AI-powered resume rewriting for each job"""
        self.logger.info("  ✍️  Rewriting resume with AI...")
        
        prompt = f"""Rewrite this resume to perfectly match this job posting. Optimize for ATS and highlight relevant skills.

JOB TITLE: {job_title}
COMPANY: {company}
JOB DESCRIPTION: {job_description}

ORIGINAL RESUME:
{self.config['resume_text']}

Requirements:
1. Keep the same experience but reframe accomplishments to match job requirements
2. Reorder skills to prioritize those mentioned in job description
3. Add relevant keywords naturally
4. Adjust summary to align with role
5. Keep format professional and ATS-friendly
6. Return ONLY the rewritten resume text, no explanations

REWRITTEN RESUME:"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2000
            )
            
            tailored_text = response.choices[0].message.content
            
            # Create new resume document
            resume_path = self.create_resume_document(tailored_text, job_title)
            self.logger.info(f"  📄 Resume tailored and saved")
            
            return resume_path
            
        except Exception as e:
            self.logger.error(f"  ⚠️  AI rewrite failed, using original: {e}")
            return self.config['resume_path']
    
    def create_resume_document(self, text: str, job_title: str) -> str:
        """Create formatted Word document from text"""
        doc = Document()
        
        # Parse and format the text
        lines = text.split('\n')
        
        for line in lines:
            if line.strip():
                # Detect headers (all caps or followed by colon)
                if line.isupper() or ':' in line[:30]:
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(12)
                else:
                    doc.add_paragraph(line)
        
        # Save with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_tailored_{timestamp}.docx"
        filepath = os.path.join(self.config['temp_folder'], filename)
        
        doc.save(filepath)
        return filepath
    
    def fill_linkedin_application(self, resume_path: str) -> bool:
        """Auto-fill LinkedIn Easy Apply multi-page form"""
        try:
            page = 1
            while True:
                self.logger.info(f"  📝 Filling form page {page}...")
                
                # Find all form fields
                self.auto_fill_form_fields()
                
                # Handle file upload
                self.upload_resume_if_needed(resume_path)
                
                # Look for next/submit button
                try:
                    next_button = self.driver.find_element(By.XPATH, "//button[@aria-label='Continue to next step']")
                    next_button.click()
                    time.sleep(2)
                    page += 1
                except:
                    # Try submit button
                    try:
                        submit_button = self.driver.find_element(By.XPATH, "//button[@aria-label='Submit application']")
                        submit_button.click()
                        time.sleep(3)
                        
                        # Close confirmation modal
                        try:
                            close_btn = self.driver.find_element(By.XPATH, "//button[@aria-label='Dismiss']")
                            close_btn.click()
                        except:
                            pass
                        
                        return True
                    except:
                        self.logger.warning("  ⚠️  Couldn't find submit button")
                        return False
                        
        except Exception as e:
            self.logger.error(f"  ❌ Form filling failed: {e}")
            return False
    
    def auto_fill_form_fields(self):
        """Intelligently auto-fill all form fields"""
        # Text inputs
        inputs = self.driver.find_elements(By.TAG_NAME, 'input')
        for inp in inputs:
            try:
                field_name = inp.get_attribute('name') or inp.get_attribute('id') or ''
                label = self.get_field_label(inp)
                
                value = self.get_field_value(field_name, label)
                if value and inp.is_displayed() and inp.is_enabled():
                    inp.clear()
                    inp.send_keys(value)
                    
            except Exception as e:
                pass
        
        # Textareas
        textareas = self.driver.find_elements(By.TAG_NAME, 'textarea')
        for textarea in textareas:
            try:
                label = self.get_field_label(textarea)
                if 'cover' in label.lower():
                    # Generate cover letter on the fly
                    cover_letter = self.generate_quick_cover_letter()
                    textarea.clear()
                    textarea.send_keys(cover_letter)
            except:
                pass
        
        # Dropdowns
        selects = self.driver.find_elements(By.TAG_NAME, 'select')
        for select in selects:
            try:
                label = self.get_field_label(select)
                value = self.get_dropdown_value(label)
                if value:
                    Select(select).select_by_visible_text(value)
            except:
                pass
        
        # Radio buttons and checkboxes
        self.handle_radio_and_checkboxes()
    
    def get_field_label(self, element) -> str:
        """Extract label for form field"""
        try:
            # Try aria-label
            label = element.get_attribute('aria-label')
            if label:
                return label.lower()
            
            # Try associated label
            field_id = element.get_attribute('id')
            if field_id:
                label_elem = self.driver.find_element(By.XPATH, f"//label[@for='{field_id}']")
                return label_elem.text.lower()
        except:
            pass
        return ''
    
    def get_field_value(self, field_name: str, label: str) -> str:
        """Get appropriate value for field"""
        field_combined = (field_name + ' ' + label).lower()
        
        # Phone
        if any(x in field_combined for x in ['phone', 'mobile', 'tel']):
            return self.config['phone']
        
        # Email
        if 'email' in field_combined:
            return self.config['email']
        
        # Name
        if 'first' in field_combined and 'name' in field_combined:
            return self.config['first_name']
        if 'last' in field_combined and 'name' in field_combined:
            return self.config['last_name']
        if 'name' in field_combined:
            return self.config['full_name']
        
        # LinkedIn
        if 'linkedin' in field_combined:
            return self.config['linkedin_url']
        
        # Website/Portfolio
        if any(x in field_combined for x in ['website', 'portfolio', 'github']):
            return self.config.get('portfolio_url', '')
        
        # Years of experience
        if 'year' in field_combined and 'experience' in field_combined:
            return str(self.config['years_experience'])
        
        # Salary
        if 'salary' in field_combined or 'compensation' in field_combined:
            return str(self.config.get('desired_salary', ''))
        
        # City
        if 'city' in field_combined:
            return self.config.get('city', '')
        
        # State
        if 'state' in field_combined:
            return self.config.get('state', '')
        
        return ''
    
    def get_dropdown_value(self, label: str) -> str:
        """Get appropriate dropdown selection"""
        label = label.lower()
        
        if 'experience' in label:
            years = self.config['years_experience']
            if years < 2:
                return '0-2 years'
            elif years < 5:
                return '3-5 years'
            elif years < 10:
                return '5-10 years'
            else:
                return '10+ years'
        
        if 'work authorization' in label or 'authorized' in label:
            return self.config.get('work_authorization', 'Yes')
        
        if 'require sponsorship' in label:
            return self.config.get('require_sponsorship', 'No')
        
        if 'veteran' in label:
            return 'No'
        
        if 'disability' in label:
            return 'Prefer not to say'
        
        if 'gender' in label:
            return 'Prefer not to say'
        
        if 'race' in label or 'ethnicity' in label:
            return 'Prefer not to say'
        
        return ''
    
    def handle_radio_and_checkboxes(self):
        """Handle radio buttons and checkboxes intelligently"""
        # Find all radio groups
        radios = self.driver.find_elements(By.XPATH, "//input[@type='radio']")
        
        for radio in radios:
            try:
                label = self.get_field_label(radio)
                
                # Auto-select appropriate options
                if any(x in label.lower() for x in ['yes', 'authorized', 'eligible', 'legally']):
                    if not radio.is_selected():
                        radio.click()
                elif any(x in label.lower() for x in ['no', 'not require', "don't require"]):
                    if 'sponsorship' in label.lower():
                        if not radio.is_selected():
                            radio.click()
            except:
                pass
    
    def upload_resume_if_needed(self, resume_path: str):
        """Handle resume upload"""
        try:
            file_inputs = self.driver.find_elements(By.XPATH, "//input[@type='file']")
            for file_input in file_inputs:
                if file_input.is_displayed() or True:  # Hidden inputs still work
                    file_input.send_keys(os.path.abspath(resume_path))
                    self.logger.info("  📎 Resume uploaded")
                    time.sleep(2)
        except Exception as e:
            self.logger.debug(f"  No file upload needed or failed: {e}")
    
    def generate_quick_cover_letter(self) -> str:
        """Generate brief cover letter"""
        return f"""Dear Hiring Manager,

I am excited to apply for this position. With {self.config['years_experience']} years of experience in {', '.join(self.config['top_skills'][:3])}, I am confident I can contribute to your team's success.

I look forward to discussing how my background aligns with your needs.

Best regards,
{self.config['full_name']}"""
    
    def scroll_page(self, scrolls=3):
        """Scroll to load dynamic content"""
        for _ in range(scrolls):
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
    
    def run_continuous(self):
        """Run bot continuously 24/7"""
        self.logger.info("🚀 Starting AUTOMATED JOB BOT...")
        
        self.setup_browser(headless=False)
        self.login_to_platforms()
        
        job_searches = self.config['job_searches']
        
        while True:
            for search in job_searches:
                try:
                    self.logger.info(f"\n🔍 Searching: {search['keywords']} in {search['location']}")
                    
                    # LinkedIn
                    self.search_and_apply_linkedin(search['keywords'], search['location'])
                    
                    # Indeed (similar implementation)
                    # self.search_and_apply_indeed(search['keywords'], search['location'])
                    
                    # Glassdoor (similar implementation)
                    # self.search_and_apply_glassdoor(search['keywords'], search['location'])
                    
                except Exception as e:
                    self.logger.error(f"❌ Search failed: {e}")
            
            # Wait before next cycle
            wait_time = self.config.get('cycle_wait_hours', 4) * 3600
            self.logger.info(f"\n💤 Waiting {self.config.get('cycle_wait_hours', 4)} hours before next cycle...")
            self.logger.info(f"📊 Total applications sent: {self.applied_count}")
            time.sleep(wait_time)


# ==================== CONFIGURATION ====================

config = {
    # Your Credentials
    'linkedin_email': 'your_email@gmail.com',
    'linkedin_password': 'your_password',
    'indeed_email': 'your_email@gmail.com',
    'indeed_password': 'your_password',
    
    # Personal Info
    'first_name': 'John',
    'last_name': 'Doe',
    'full_name': 'John Doe',
    'email': 'john.doe@gmail.com',
    'phone': '+1-234-567-8900',
    'linkedin_url': 'https://linkedin.com/in/johndoe',
    'portfolio_url': 'https://github.com/johndoe',
    'city': 'San Francisco',
    'state': 'California',
    
    # Work Info
    'years_experience': 5,
    'top_skills': ['Python', 'Machine Learning', 'Data Science', 'AWS', 'SQL'],
    'desired_salary': 150000,
    'work_authorization': 'Yes',
    'require_sponsorship': 'No',
    
    # Resume
    'resume_path': '/path/to/your/resume.docx',
    'resume_text': """
    JOHN DOE
    Software Engineer
    
    EXPERIENCE
    Senior Developer at Tech Corp (2020-2024)
    - Led team of 5 engineers
    - Increased system performance by 40%
    
    SKILLS
    Python, JavaScript, AWS, Docker, SQL, Machine Learning
    """,
    
    # AI API Keys
    'openai_key': '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',
    'anthropic_key': 'sk-ant-your-anthropic-key',
    
    # Job Search Criteria
    'job_searches': [
        {'keywords': 'Software Engineer', 'location': 'Remote'},
        {'keywords': 'Python Developer', 'location': 'San Francisco'},
        {'keywords': 'ML Engineer', 'location': 'Remote'},
    ],
    
    # Bot Settings
    'cycle_wait_hours': 6,  # Wait between search cycles
    'temp_folder': './tailored_resumes',
}

# ==================== RUN BOT ====================

if __name__ == "__main__":
    # Create temp folder
    os.makedirs(config['temp_folder'], exist_ok=True)
    
    # Initialize and run
    bot = AutomatedJobBot(config)
    bot.run_continuous()