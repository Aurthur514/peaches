"""
Task 3 - GLR Pipeline (Streamlit)

Streamlit app that accepts a .docx insurance template and multiple PDF photo
reports, extracts text from the PDFs, asks an LLM to map template fields to
values found in the reports, and returns a filled .docx for download.

Usage (local):
  .venv\Scripts\Activate.ps1
  pip install -r task_3_requirements.txt
  streamlit run task_3_glr_streamlit.py

Environment variables:
  - OPENROUTER_API_KEY : (optional) API key for OpenRouter. If set, will try
    OpenRouter first.
  - OPENAI_API_KEY : (optional) fallback to OpenAI ChatCompletions API.

Notes:
  - Template placeholders are expected as {{field_name}} in the .docx. If the
    template does not contain explicit placeholders, the app will ask the LLM
    to propose fields based on the template content.
"""

from io import BytesIO
import os
import re
import json
import requests
from typing import List, Dict, Tuple

import streamlit as st
try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    Document = None
    HAVE_DOCX = False
try:
    from PyPDF2 import PdfReader
    HAVE_PYPDF2 = True
except Exception:
    PdfReader = None
    HAVE_PYPDF2 = False
import sys
import platform
import importlib
import importlib.metadata as importlib_metadata
import io
from task_3_internal_llm import propose_fields_from_template, map_fields_from_reports



# ------------------ Utility functions ------------------

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}\s]+)\s*\}\}")


def extract_text_from_pdfs(uploaded_files: List[BytesIO]) -> str:
    texts = []
    for f in uploaded_files:
        try:
            # read bytes to allow multiple passes (text extraction and OCR)
            try:
                f.seek(0)
            except Exception:
                pass
            pdf_bytes = f.read()
            if not pdf_bytes:
                st.warning("Empty PDF file uploaded; skipping.")
                continue

            # First try text extraction via PyPDF2
            pages_text = []
            try:
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for p in reader.pages:
                    try:
                        pages_text.append(p.extract_text() or "")
                    except Exception:
                        pages_text.append("")
            except Exception as e:
                st.warning(f"PyPDF2 failed to read PDF: {e}")

            combined = "\n\n".join(pages_text).strip()

            # If extracted text is empty or very short, attempt OCR fallback
            if len(combined) < 100:
                # Lazy import OCR libraries
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    from PIL import Image
                except Exception as e:
                    st.warning("OCR libraries not available (install pytesseract, pdf2image, pillow). Skipping OCR.")
                    texts.append(combined)
                    continue

                st.info("No (or little) searchable text found — running OCR on PDF pages (this may be slow).")
                try:
                    poppler_path = os.environ.get('POPPLER_PATH') or None
                    images = convert_from_bytes(pdf_bytes, dpi=300, fmt='jpeg', poppler_path=poppler_path)
                    ocr_pages = []
                    for img in images:
                        try:
                            txt = pytesseract.image_to_string(img)
                            ocr_pages.append(txt or "")
                        except Exception:
                            ocr_pages.append("")
                    combined = "\n\n".join(ocr_pages).strip()
                except Exception as e:
                    st.warning(f"pdf2image / pytesseract OCR failed: {e}")

            texts.append(combined)
        except Exception as e:
            st.warning(f"Failed to process a PDF: {e}")
    return "\n\n".join(texts)


def extract_text_from_docx(file_like: BytesIO) -> Tuple[str, Document]:
    # Return (plain_text, Document object)
    doc = Document(file_like)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts), doc


def find_placeholders(doc_text: str) -> List[str]:
    return list(dict.fromkeys(PLACEHOLDER_RE.findall(doc_text)))


# ------------------ LLM integration ------------------


def call_llm_chat(messages: List[Dict], model: str = None, timeout: int = 60) -> str:
    """Try OpenRouter first if OPENROUTER_API_KEY present; fall back to OpenAI.

    Messages should follow Chat Completions format: [{'role': 'user'|'system', 'content': '...'}]
    Returns the assistant text.
    """
    # Support multiple providers: GEMINI (Google), OpenRouter, OpenAI
    gemini_key = os.environ.get("GEMINI_API_KEY")
    openrouter_key = os.environ.get("OPENROUTER_API_KEY")
    openai_key = os.environ.get("OPENAI_API_KEY")

    headers = None
    payload = None

    # Prefer OpenRouter/OpenAI when available (more stable for bearer tokens), then fall back to Gemini
    # 1) Try OpenRouter (if key)
    if openrouter_key:
        url = "https://api.openrouter.ai/v1/chat/completions"
        headers = {"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"}
        payload = {
            "model": model or "gpt-4o-mini",
            "messages": messages,
            "temperature": 0.0,
            "max_tokens": 1500,
        }
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=timeout)
            r.raise_for_status()
            j = r.json()
            return j['choices'][0]['message']['content']
        except Exception as e:
            st.warning(f"OpenRouter request failed: {e}")
            try:
                print(f"OpenRouter request failed: {e}")
            except Exception:
                pass

    # 2) Try OpenAI (if key)
    if openai_key:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {openai_key}", "Content-Type": "application/json"}
        payload = {
            "model": model or "gpt-4o-mini",
            "messages": messages,
            "temperature": 0.0,
            "max_tokens": 1500,
        }
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=timeout)
            r.raise_for_status()
            j = r.json()
            return j['choices'][0]['message']['content']
        except Exception as e:
            st.warning(f"OpenAI request failed: {e}")
            try:
                print(f"OpenAI request failed: {e}")
            except Exception:
                pass

    # 3) Try Google Generative Language (Gemini) if key provided
    if gemini_key:
        try:
            # Convert chat-like messages to a single prompt for text generation
            prompt = []
            for m in messages:
                role = m.get('role', '')
                content = m.get('content', '')
                prompt.append(f"[{role}] {content}")
            prompt_text = "\n\n".join(prompt)

            url = f"https://generativelanguage.googleapis.com/v1beta2/models/text-bison-001:generateText?key={gemini_key}"
            payload = {
                "prompt": {"text": prompt_text},
                "temperature": 0.0,
                "maxOutputTokens": 1500,
            }
            r = requests.post(url, json=payload, timeout=timeout)
            r.raise_for_status()
            j = r.json()
            # Response contains candidates -> output
            if 'candidates' in j and len(j['candidates']) > 0:
                return j['candidates'][0].get('output', '')
        except Exception as e:
            st.warning(f"Gemini request failed: {e} — check GEMINI_API_KEY and model/endpoint")
            try:
                print(f"Gemini request failed: {e}")
            except Exception:
                pass

    print("No LLM providers succeeded. Available env keys:", {"GEMINI": bool(gemini_key), "OPENROUTER": bool(openrouter_key), "OPENAI": bool(openai_key)})
    raise RuntimeError("No LLM API keys available or all providers failed (check console output)")


def ask_llm_for_fields(template_text: str) -> List[str]:
    system = (
        "You are an assistant that extracts the list of fields that should be filled in an insurance\n"
        "template document. Return a JSON array of short field names (no spaces) suitable as placeholders.\n"
        "Example output: [\"policy_number\", \"insured_name\"]\n"
        "Do not include any extra text."
    )
    user = f"Template text:\n\n{template_text}\n\nReturn the list of fields as described."
    resp = call_llm_chat([{"role": "system", "content": system}, {"role": "user", "content": user}])
    # Try to parse JSON out of the response
    try:
        out = json.loads(resp)
        if isinstance(out, list):
            return [str(x).strip() for x in out]
    except Exception:
        # fallback: try to find words separated by commas/newlines
        candidates = re.findall(r"[A-Za-z0-9_\-]{2,}", resp)
        return list(dict.fromkeys(candidates))


def ask_llm_to_map(fields: List[str], reports_text: str) -> Dict[str, str]:
    system = (
        "You are an assistant that extracts values for given template fields from a set of photo report texts.\n"
        "Given a JSON array of field names and the combined report text, return a JSON object mapping\n"
        "each field name to the best matching value found in the reports. If a value is not present, use empty string.\n"
        "Output only valid JSON."
    )
    user = (
        f"Fields: {json.dumps(fields)}\n\n"
        f"Reports text:\n{reports_text[:12000]}"  # truncate to keep payload reasonable
    )
    resp = call_llm_chat([{"role": "system", "content": system}, {"role": "user", "content": user}], model="gpt-4o-mini")
    try:
        mapping = json.loads(resp)
        if isinstance(mapping, dict):
            return {k: (v if isinstance(v, str) else str(v)) for k, v in mapping.items()}
    except Exception:
        # Try to extract a JSON-looking block from the response
        jmatch = re.search(r"\{[\s\S]*\}", resp)
        if jmatch:
            try:
                mapping = json.loads(jmatch.group(0))
                return {k: (v if isinstance(v, str) else str(v)) for k, v in mapping.items()}
            except Exception:
                pass
    # If we got here, fallback: return empty values
    return {f: "" for f in fields}


# ------------------ Docx filling ------------------


def fill_docx_template(doc: Document, mapping: Dict[str, str]) -> Document:
    """Replace placeholders like {{field}} in paragraphs and table cells with mapping values."""
    def replace_in_run_text(text: str) -> str:
        def _rep(m):
            key = m.group(1)
            return mapping.get(key, m.group(0))
        return PLACEHOLDER_RE.sub(_rep, text)

    # Paragraphs
    for p in doc.paragraphs:
        if PLACEHOLDER_RE.search(p.text):
            # Runs may split placeholders; rebuild paragraph text and re-assign.
            full = ''.join([r.text for r in p.runs])
            replaced = replace_in_run_text(full)
            # Clear runs and set a single run with replaced text
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if PLACEHOLDER_RE.search(cell.text):
                    text = '\n'.join([p.text for p in cell.paragraphs])
                    replaced = replace_in_run_text(text)
                    # clear existing paragraphs
                    for p in cell.paragraphs:
                        p.text = ''
                    cell.paragraphs[0].add_run(replaced)

    return doc


# ------------------ Streamlit app ------------------


def main():
    st.set_page_config(page_title="Task 3 - GLR Pipeline", layout="wide")
    st.title("Task 3 — GLR Pipeline (Insurance Template Filling)")

    # Diagnostics panel (non-sensitive): show Python and package info, and whether key env vars exist
    with st.sidebar.expander("Diagnostics", expanded=False):
        if st.button("Refresh diagnostics"):
            pass
        def _get_pkg_version(pkg_name, module_name=None):
            # Try metadata first, then module __version__, else not installed
            try:
                return importlib_metadata.version(pkg_name)
            except Exception:
                try:
                    if module_name:
                        mod = importlib.import_module(module_name)
                    else:
                        mod = importlib.import_module(pkg_name)
                    return getattr(mod, "__version__", "unknown")
                except Exception:
                    return "not installed"

        pkgs = {
            "python": platform.python_version(),
            "streamlit": _get_pkg_version("streamlit"),
            "python-docx": _get_pkg_version("python-docx", "docx"),
            "PyPDF2": _get_pkg_version("PyPDF2", "PyPDF2"),
            "pdf2image": _get_pkg_version("pdf2image"),
            "pytesseract": _get_pkg_version("pytesseract"),
            "Pillow": _get_pkg_version("Pillow", "PIL"),
        }

        st.markdown("**Runtime**")
        st.write({"python_version": platform.python_version(), "executable": sys.executable})
        st.markdown("**Key package versions**")
        st.table([{"package": k, "version": v} for k, v in pkgs.items()])
        st.markdown("**Environment keys present (values hidden)**")
        keys = ["GEMINI_API_KEY", "OPENROUTER_API_KEY", "OPENAI_API_KEY", "POPPLER_PATH", "TESSERACT_CMD"]
        st.write({k: bool(os.environ.get(k)) for k in keys})

    if not HAVE_DOCX or not HAVE_PYPDF2:
        missing = []
        if not HAVE_DOCX:
            missing.append('python-docx')
        if not HAVE_PYPDF2:
            missing.append('PyPDF2')
        st.error(
            "Missing packages: `" + ", ".join(missing) + "`.\n"
            "Install them and redeploy.\n"
            "On Streamlit Cloud add the missing packages to your repository `requirements.txt` or in the app's package settings.\n"
            "Locally run: `pip install " + " ".join(missing) + "` and then `streamlit run task_3_glr_streamlit.py`."
        )
        return

    st.markdown("Upload a `.docx` template (placeholders like `{{field}}`) and one or more PDF photo reports.")

    # Provider selector: internal heuristic (no external LLM) or auto (use configured providers)
    provider_choice = st.radio("Provider", ["Internal (heuristic, no external LLMs)", "Auto (use configured LLM providers)"], index=0)

    with st.form("upload_form"):
        template_file = st.file_uploader("Insurance template (.docx)", type=["docx"], accept_multiple_files=False)
        pdf_files = st.file_uploader("Photo report PDFs (multiple)", type=["pdf"], accept_multiple_files=True)
        min_area = st.number_input("Min area for change detection (unused here)", value=500)
        submit = st.form_submit_button("Process")

    if not submit:
        st.info("Upload files and click Process.")
        return

    if not template_file:
        st.error("Please upload a .docx template.")
        return
    if not pdf_files:
        st.error("Please upload at least one PDF report.")
        return

    st.info("Extracting text from files...")
    # Read template text
    template_bytes = template_file.read()
    template_text, template_doc = extract_text_from_docx(BytesIO(template_bytes))

    # Extract PDFs text
    pdf_text = extract_text_from_pdfs([BytesIO(p.read()) for p in pdf_files])

    # Find placeholders
    placeholders = find_placeholders(template_text)
    if placeholders:
        st.success(f"Found placeholders in template: {placeholders}")
    else:
        if provider_choice.startswith('Internal'):
            st.info("No {{placeholders}} found in template. Using internal heuristic to propose fields...")
            placeholders = propose_fields_from_template(template_text)
            st.success(f"Internal proposed fields: {placeholders}")
        else:
            st.warning("No {{placeholders}} found in template. Asking LLM to propose fields...")
            try:
                placeholders = ask_llm_for_fields(template_text)
                st.success(f"LLM proposed fields: {placeholders}")
            except Exception as e:
                st.error(f"LLM failed to propose fields: {e}")
                return

    if provider_choice.startswith('Internal'):
        st.info("Using internal heuristic to map fields (no external LLM calls).")
        mapping = map_fields_from_reports(placeholders, pdf_text)
    else:
        st.info("Asking LLM to map fields to values from the reports (this requires an API key)...")
        try:
            mapping = ask_llm_to_map(placeholders, pdf_text)
        except Exception as e:
            st.error(f"LLM mapping failed: {e}")
            return

    st.subheader("Extracted mapping")
    st.json(mapping)

    # Allow user to edit mapping manually
    edited = st.text_area("Edit JSON mapping before applying", value=json.dumps(mapping, indent=2), height=200)
    try:
        mapping_final = json.loads(edited)
    except Exception as e:
        st.error(f"Invalid JSON: {e}")
        return

    st.info("Filling template...")
    filled_doc = fill_docx_template(Document(BytesIO(template_bytes)), mapping_final)

    # Save to bytes
    out_stream = BytesIO()
    filled_doc.save(out_stream)
    out_stream.seek(0)

    st.success("Template filled — download below")
    st.download_button("Download filled .docx", data=out_stream.getvalue(), file_name="filled_template.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Also offer the mapping as a JSON download
    st.download_button("Download mapping JSON", data=json.dumps(mapping_final, indent=2), file_name="mapping.json", mime='application/json')


if __name__ == '__main__':
    main()
